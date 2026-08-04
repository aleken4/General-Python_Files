import os
import sys
import glob
import time
from collections import defaultdict
import numpy as np
import face_recognition
import cv2
from moviepy.editor import VideoFileClip
import webrtcvad
import wave
import contextlib
import audioop
from pydub import AudioSegment
from docx import Document
from docx.shared import Inches
from dataclasses import dataclass

# ---------- Config ----------
DATASET_DIR = "dataset"       # dataset/{name}/*.jpg
TEMP_AUDIO = "temp_audio.wav"
FRAME_INTERVAL = 0.5         # seconds between frames to analyze
VAD_FRAME_MS = 30
VAD_MODE = 2                # 0-3, 3 more aggressive
OUTPUT_DOCX = "speakers.docx"
# ----------------------------

@dataclass
class FaceMatch:
    name: str
    distance: float
    bbox: tuple  # top, right, bottom, left

def load_known_faces(dataset_dir=DATASET_DIR):
    known_encodings = []
    known_names = []
    for person_dir in os.listdir(dataset_dir):
        person_path = os.path.join(dataset_dir, person_dir)
        if not os.path.isdir(person_path):
            continue
        for img_path in glob.glob(os.path.join(person_path, "*")):
            try:
                img = face_recognition.load_image_file(img_path)
                encs = face_recognition.face_encodings(img)
                if len(encs) == 0:
                    continue
                known_encodings.append(encs[0])
                known_names.append(person_dir)
            except Exception as e:
                print("Error loading", img_path, e)
    print(f"Loaded {len(known_encodings)} face encodings.")
    return known_encodings, known_names

# Extract audio from video using moviepy and save as mono 16-bit 16kHz wav
def extract_audio(video_path, out_wav=TEMP_AUDIO):
    clip = VideoFileClip(video_path)
    clip.audio.write_audiofile("tmp_audio.mp3", verbose=False, logger=None)
    # convert to mono 16kHz wav
    sound = AudioSegment.from_file("tmp_audio.mp3")
    sound = sound.set_frame_rate(16000).set_channels(1).set_sample_width(2)
    sound.export(out_wav, format="wav")
    os.remove("tmp_audio.mp3")
    return out_wav, clip.duration

# Simple VAD-based voiced segments (returns list of (start, end) in seconds)
def vad_segments(wav_path, frame_ms=VAD_FRAME_MS, mode=VAD_MODE):
    vad = webrtcvad.Vad(mode)
    segments = []
    with wave.open(wav_path, "rb") as wf:
        rate = wf.getframerate()
        channels = wf.getnchannels()
        width = wf.getsampwidth()
        assert channels == 1, "audio must be mono"
        bytes_per_frame = int(rate * (frame_ms / 1000.0) * width)
        timestamp = 0.0
        voiced = False
        seg_start = None
        while True:
            frame = wf.readframes(int(rate*frame_ms/1000.0))
            if len(frame) == 0:
                if voiced:
                    segments.append((seg_start, timestamp))
                break
            is_speech = vad.is_speech(frame, sample_rate=rate)
            if is_speech and not voiced:
                voiced = True
                seg_start = timestamp
            elif not is_speech and voiced:
                voiced = False
                segments.append((seg_start, timestamp))
            timestamp += frame_ms/1000.0
    # merge short gaps
    merged = []
    for s,e in segments:
        if not merged:
            merged.append([s,e])
        else:
            if s - merged[-1][1] < 0.25:
                merged[-1][1] = e
            else:
                merged.append([s,e])
    return [(a,b) for a,b in merged]

# Analyze frames at regular intervals, return for each timestamp the list of face matches
def analyze_frames(video_path, known_encodings, known_names, frame_interval=FRAME_INTERVAL):
    cap = cv2.VideoCapture(video_path)
    fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
    duration = total_frames / fps if fps else 0
    timestamps = np.arange(0, duration, frame_interval)
    results = {}  # timestamp -> list of FaceMatch
    for t in timestamps:
        frame_no = int(t * fps)
        cap.set(cv2.CAP_PROP_POS_FRAMES, frame_no)
        ret, frame = cap.read()
        if not ret:
            continue
        rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        boxes = face_recognition.face_locations(rgb, model="hog")
        encs = face_recognition.face_encodings(rgb, boxes)
        matches = []
        for enc, box in zip(encs, boxes):
            if len(known_encodings) == 0:
                name = "Unknown"
                dist = 1.0
            else:
                dists = face_recognition.face_distance(known_encodings, enc)
                idx = np.argmin(dists)
                dist = float(dists[idx])
                # threshold ~0.5-0.6; adjust as needed
                name = known_names[idx] if dist < 0.55 else "Unknown"
            matches.append(FaceMatch(name=name, distance=dist, bbox=box))
        results[round(t,2)] = matches
    cap.release()
    return results

# Align voiced segments with frame matches to decide speaker name
def detect_speakers(voiced_segments, frame_matches):
    speaker_events = []  # (start, end, name)
    for (s,e) in voiced_segments:
        # examine frames in this interval
        ts = [t for t in frame_matches.keys() if t >= s and t <= e]
        if not ts:
            continue
        # accumulate counts weighted by face area
        score = defaultdict(float)
        for t in ts:
            matches = frame_matches[t]
            for m in matches:
                top,right,bottom,left = m.bbox
                area = max(1,(bottom-top)*(right-left))
                weight = 1.0/ (1.0 + m.distance)  # closer distance -> higher weight
                score[m.name] += area * weight
        if not score:
            name = "Unknown"
        else:
            name = max(score.items(), key=lambda x: x[1])[0]
        speaker_events.append((s,e,name))
    return speaker_events

# Generate docx
def write_docx(events, out_path=OUTPUT_DOCX):
    doc = Document()
    doc.add_heading('Detected Speakers', level=1)
    for s,e,name in events:
        p = doc.add_paragraph()
        p.add_run(f"{s:.2f} - {e:.2f}: ").bold = True
        p.add_run(name)
    doc.save(out_path)
    print("Saved", out_path)

# Main routine
def main(video_path):
    known_encodings, known_names = load_known_faces()
    if not os.path.exists(video_path):
        print("Video not found")
        return
    wav, duration = extract_audio(video_path)
    vad_segs = vad_segments(wav)
    frames = analyze_frames(video_path, known_encodings, known_names)
    events = detect_speakers(vad_segs, frames)
    write_docx(events)
    os.remove(wav)
    print("Done.")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python talk_recognizer.py path_to_video.mp4")
        sys.exit(1)
    main(sys.argv[1])